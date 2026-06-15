"""
File parsing service for document ingestion.
Supports: PDF, DOCX, TXT, CSV, XLSX, and other text-based formats.
"""

import logging
import os
import re
from typing import Any

from app.config import get_settings

logger = logging.getLogger(__name__)


class FileParser:
    """
    Parse uploaded files into text chunks for RAG ingestion.
    
    Supports:
    - PDF (via PyMuPDF)
    - DOCX (via python-docx)
    - TXT, MD, JSON, code files (plain text)
    - CSV, XLSX (via pandas)
    
    Chunking strategy:
    - Split by paragraphs/sections
    - Target ~500 tokens per chunk with 50-token overlap
    - Preserve section headings as metadata
    """

    def __init__(self):
        """Initialize with settings from configuration."""
        settings = get_settings()
        self.chunk_size_words = settings.chunk_size_words
        self.chunk_overlap_words = settings.chunk_overlap_words

    async def parse_file(
        self, file_path: str, filename: str
    ) -> list[dict[str, Any]]:
        """
        Parse a file and return a list of text chunks.
        
        Returns:
            List of dicts with: content, source_type, metadata
        """
        ext = os.path.splitext(filename)[1].lower()

        try:
            if ext == ".pdf":
                text = await self._parse_pdf(file_path)
            elif ext == ".docx":
                text = await self._parse_docx(file_path)
            elif ext in (".csv",):
                text = await self._parse_csv(file_path)
            elif ext in (".xlsx", ".xls"):
                text = await self._parse_excel(file_path)
            elif ext in (".txt", ".md", ".json", ".py", ".js", ".java", ".cpp",
                         ".c", ".html", ".css", ".xml", ".yaml", ".yml", ".ini",
                         ".cfg", ".log", ".rst", ".tex"):
                text = await self._parse_text(file_path)
            else:
                # Try as plain text
                text = await self._parse_text(file_path)

            if not text.strip():
                logger.warning(f"No text extracted from {filename}")
                return []

            # Sanitize text (remove null bytes, non-printable chars)
            text = self._sanitize_text(text)

            # Chunk the text
            chunks = self._chunk_text(text, filename, ext)
            logger.info(f"Parsed {filename}: {len(text)} chars → {len(chunks)} chunks")
            return chunks

        except Exception as e:
            logger.error(f"Failed to parse {filename}: {e}", exc_info=True)
            raise

    async def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF using PyMuPDF."""
        import asyncio

        def _extract():
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            text_parts = []
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            doc.close()
            return "\n\n".join(text_parts)

        return await asyncio.to_thread(_extract)

    async def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX using python-docx."""
        import asyncio

        def _extract():
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Preserve heading structure
                    if paragraph.style.name.startswith("Heading"):
                        level = paragraph.style.name.replace("Heading ", "")
                        text_parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {paragraph.text}")
                    else:
                        text_parts.append(paragraph.text)

            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)

        return await asyncio.to_thread(_extract)

    async def _parse_csv(self, file_path: str) -> str:
        """Parse CSV using pandas."""
        import asyncio

        def _extract():
            import pandas as pd
            df = pd.read_csv(file_path, nrows=10000)  # Limit rows
            # Convert to readable text
            text_parts = []
            text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
            text_parts.append(f"Rows: {len(df)}")
            text_parts.append("")
            # Convert rows to text
            for _, row in df.iterrows():
                row_text = ". ".join(
                    f"{col}: {val}" for col, val in row.items()
                    if pd.notna(val) and str(val).strip()
                )
                if row_text:
                    text_parts.append(row_text)
            return "\n".join(text_parts)

        return await asyncio.to_thread(_extract)

    async def _parse_excel(self, file_path: str) -> str:
        """Parse Excel using pandas."""
        import asyncio

        def _extract():
            import pandas as pd
            # Read all sheets
            sheets = pd.read_excel(file_path, sheet_name=None, nrows=10000)
            text_parts = []
            for sheet_name, df in sheets.items():
                text_parts.append(f"## Sheet: {sheet_name}")
                text_parts.append(f"Columns: {', '.join(df.columns.tolist())}")
                for _, row in df.iterrows():
                    row_text = ". ".join(
                        f"{col}: {val}" for col, val in row.items()
                        if pd.notna(val) and str(val).strip()
                    )
                    if row_text:
                        text_parts.append(row_text)
                text_parts.append("")
            return "\n".join(text_parts)

        return await asyncio.to_thread(_extract)

    async def _parse_text(self, file_path: str) -> str:
        """Parse plain text files."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def _split_into_topics(self, text: str) -> list[tuple[str, str]]:
        """
        Split text into topics based on common header patterns.
        Returns a list of tuples: (topic_title, topic_content)
        """
        # Regex to match Markdown headers, Chapters, Sections, Topics, or Excel sheets
        header_pattern = re.compile(
            r"^(?:#{1,6}\s+|Chapter\s+\d+|Section\s+[\d\.]+|Topic\s+\d+|## Sheet:).*$",
            re.IGNORECASE | re.MULTILINE
        )
        
        topics = []
        last_idx = 0
        current_title = "Document Intro"
        
        for match in header_pattern.finditer(text):
            start = match.start()
            # Save the text before this header under the current_title
            content = text[last_idx:start].strip()
            if content:
                topics.append((current_title, content))
                
            current_title = match.group().strip()
            last_idx = match.end()
            
        # Add the last section
        content = text[last_idx:].strip()
        if content:
            topics.append((current_title, content))
            
        return topics

    def _chunk_text(
        self, text: str, filename: str, file_ext: str
    ) -> list[dict[str, Any]]:
        """
        Split text using a two-phase approach:
        1. Topic-wise chunking (by semantic headers)
        2. Sentence-wise chunking within each topic
        """
        chunks = []
        source_type = self._get_source_type(file_ext)
        
        # Phase 1: Topic-wise split
        topics = self._split_into_topics(text)
        
        # Simple sentence boundary regex (looks for .!? followed by space and Capital letter)
        sentence_end_pattern = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')
        
        # Phase 2: Sentence-wise split within each topic
        for topic_title, topic_content in topics:
            sentences = sentence_end_pattern.split(topic_content)
            
            if not sentences:
                continue
                
            current_chunk_sentences = []
            current_word_count = 0
            
            for sentence in sentences:
                words = sentence.split()
                if not words:
                    continue
                    
                current_chunk_sentences.append(sentence)
                current_word_count += len(words)
                
                if current_word_count >= self.chunk_size_words:
                    chunks.append({
                        "content": " ".join(current_chunk_sentences),
                        "source_type": source_type,
                        "metadata": {
                            "filename": filename,
                            "topic": topic_title
                        },
                    })
                    
                    # Keep the last sentence for overlap
                    current_chunk_sentences = current_chunk_sentences[-1:]
                    current_word_count = len(current_chunk_sentences[0].split())
            
            # Add remaining sentences as a final chunk
            if current_chunk_sentences and current_word_count > 0:
                chunks.append({
                    "content": " ".join(current_chunk_sentences),
                    "source_type": source_type,
                    "metadata": {
                        "filename": filename,
                        "topic": topic_title
                    },
                })
                    
        return chunks

    def _get_source_type(self, ext: str) -> str:
        """Map file extension to source type."""
        mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".doc": "doc",
            ".txt": "text",
            ".md": "markdown",
            ".csv": "csv",
            ".xlsx": "excel",
            ".xls": "excel",
            ".json": "json",
        }
        return mapping.get(ext, "document")

    @staticmethod
    def _sanitize_text(text: str) -> str:
        """Remove null bytes and non-printable characters that PostgreSQL rejects."""
        text = text.replace('\x00', '')
        text = re.sub(r'[^\x09\x0a\x0d\x20-\x7e\u00a0-\uffff]', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
